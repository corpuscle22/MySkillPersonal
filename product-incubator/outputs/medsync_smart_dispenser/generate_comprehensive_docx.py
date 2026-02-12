"""Generate comprehensive Word document for MedSync Smart Dispenser."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

def add_heading_style(doc):
    """Customize heading styles."""
    styles = doc.styles

    # Modify Heading 1
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.size = Pt(16)
    h1_font.bold = True
    h1_font.color.rgb = RGBColor(0, 51, 102)

    # Modify Heading 2
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.size = Pt(14)
    h2_font.bold = True
    h2_font.color.rgb = RGBColor(0, 102, 153)

def create_document():
    doc = Document()
    add_heading_style(doc)

    # Title
    title = doc.add_heading('MedSync Smart Dispenser', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('Product Incubator Plan')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=1)

    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    cells = [
        ('Product Name', 'MedSync Smart Dispenser'),
        ('Category', 'Connected Medical Device / Digital Health'),
        ('Target Market', 'Home healthcare, elderly care, chronic disease management'),
        ('Timeline', '12-18 months to market'),
        ('Regulatory', 'FDA Class II Medical Device (510(k))')
    ]
    for i, (label, value) in enumerate(cells):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        'The MedSync Smart Dispenser is a connected medication management system featuring '
        '20 dedicated compartments for various medications (tablets, capsules, pills). '
        'The device integrates with a mobile application for prescription management, '
        'automated scheduling, and precise dose dispensing with real-time adherence monitoring.'
    )

    # Phase 1
    doc.add_heading('Phase 1: Conceptualization & Design', level=1)

    doc.add_heading('Problem Statement', level=2)
    doc.add_paragraph('Medication non-adherence is a critical healthcare challenge:')
    bullets = [
        '50% of patients with chronic diseases do not take medications as prescribed',
        '$528 billion annual cost of medication non-adherence in the US',
        '125,000 deaths annually attributed to non-adherence',
        '30-50% of hospital readmissions linked to medication issues'
    ]
    for b in bullets:
        p = doc.add_paragraph(b, style='List Bullet')

    doc.add_heading('Target Users', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['User Type', 'Needs', 'Pain Points']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    users = [
        ('Elderly patients', 'Simple interface, loud alerts, easy loading', 'Complex regimens, forgetfulness, dexterity issues'),
        ('Chronic disease patients', 'Multi-med management, timing precision', 'Polypharmacy confusion, missed doses'),
        ('Caregivers', 'Remote monitoring, refill alerts', 'Cannot be present 24/7, anxiety'),
        ('Healthcare providers', 'Adherence data, intervention triggers', 'Limited visibility between visits')
    ]
    for i, (user, needs, pain) in enumerate(users):
        table.rows[i+1].cells[0].text = user
        table.rows[i+1].cells[1].text = needs
        table.rows[i+1].cells[2].text = pain

    doc.add_heading('Core Features (MVP)', level=2)
    features = [
        '20-Compartment Storage System - Individual sealed compartments for different medications with ~30-day supply per compartment',
        'Precision Dispensing Mechanism - Stepper motor-driven system with optical/weight sensors for pill counting',
        'Mobile Application (iOS/Android) - Prescription upload, schedule configuration, remote dispense, adherence dashboard',
        'Connectivity - Wi-Fi (primary), Bluetooth LE (backup), optional LTE cellular',
        'Alert System - On-device LEDs and audio, app push notifications, SMS/call escalation to caregivers'
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('Product Specifications', level=2)
    specs = doc.add_table(rows=8, cols=2)
    specs.style = 'Table Grid'
    spec_data = [
        ('Dimensions', '12" x 10" x 6" (305 x 254 x 152 mm)'),
        ('Weight', '3-4 lbs (1.4-1.8 kg)'),
        ('Compartments', '20 individual compartments'),
        ('Capacity per Compartment', '~60 medium tablets'),
        ('Pill Size Range', '4mm to 22mm diameter'),
        ('Power', '12V/2A AC adapter + Li-ion battery backup'),
        ('Battery Backup', '24-48 hour standby'),
        ('IP Rating', 'IPX1 (drip-proof)')
    ]
    for i, (spec, value) in enumerate(spec_data):
        specs.rows[i].cells[0].text = spec
        specs.rows[i].cells[1].text = value
        specs.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    # Phase 2
    doc.add_heading('Phase 2: Intellectual Property Assessment', level=1)

    doc.add_heading('Prior Art Analysis', level=2)
    doc.add_paragraph('Key patents identified in the medication dispensing space:')

    patents = doc.add_table(rows=6, cols=3)
    patents.style = 'Table Grid'
    patent_headers = ['Patent', 'Title', 'Relevance']
    for i, h in enumerate(patent_headers):
        patents.rows[0].cells[i].text = h
        patents.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    patent_data = [
        ('US7359765B2', 'Electronic Pill Dispenser', 'High - similar feature set'),
        ('US5752621A', 'Smart Automatic Medication Dispenser', 'Medium - older patent'),
        ('US20220096330A1', 'Smart Pill Dispenser', 'High - modern connected device'),
        ('US20220375564A1', 'Smart Medication Dispenser', 'High - directly relevant'),
        ('US20140358278A1', 'Smart Automated Pill Dispenser', 'High - monitoring features')
    ]
    for i, (patent, title, rel) in enumerate(patent_data):
        patents.rows[i+1].cells[0].text = patent
        patents.rows[i+1].cells[1].text = title
        patents.rows[i+1].cells[2].text = rel

    doc.add_heading('IP Strategy Recommendations', level=2)
    ip_items = [
        'Provisional Patent Application - File on novel UI/UX elements and loading workflow',
        'Design Patents - Protect distinctive device appearance',
        'Trade Secrets - Pill counting algorithms, scheduling optimization',
        'Trademark - "MedSync" brand registration'
    ]
    for item in ip_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()
    disclaimer = doc.add_paragraph('Disclaimer: This is a preliminary assessment, not legal advice. Consult a registered patent attorney.')
    disclaimer.runs[0].italic = True

    # Phase 3
    doc.add_heading('Phase 3: Technical Planning', level=1)

    doc.add_heading('Hardware Bill of Materials', level=2)
    bom = doc.add_table(rows=17, cols=4)
    bom.style = 'Table Grid'
    bom_headers = ['Component', 'Part Example', 'Qty', 'Est. Cost']
    for i, h in enumerate(bom_headers):
        bom.rows[0].cells[i].text = h
        bom.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    bom_data = [
        ('Microcontroller', 'ESP32-WROOM-32E', '1', '$3.50'),
        ('Stepper Motor', '28BYJ-48', '2-4', '$2.00'),
        ('Stepper Driver', 'ULN2003 or DRV8825', '2-4', '$1.50'),
        ('Servo Motor', 'SG90 or MG996R', '1-2', '$3.00'),
        ('IR Break-beam Sensor', 'ITR9608', '4-8', '$0.50'),
        ('Load Cell + HX711', '5kg load cell', '1', '$5.00'),
        ('OLED Display', 'SSD1306 128x64', '1', '$4.00'),
        ('Buzzer/Speaker', 'Piezo', '1', '$1.50'),
        ('RGB LED Strip', 'WS2812B', '1', '$3.00'),
        ('RTC Module', 'DS3231', '1', '$2.00'),
        ('Li-ion Battery', '18650 + BMS', '1', '$8.00'),
        ('Power Supply', '12V 2A adapter', '1', '$5.00'),
        ('PCB', 'Custom design', '1', '$15.00'),
        ('Enclosure', 'Custom molded', '1', '$25.00'),
        ('Compartments', 'Custom molded', '20', '$1.00 ea'),
        ('Misc Hardware', 'Screws, wires, etc.', '-', '$10.00')
    ]
    for i, (comp, part, qty, cost) in enumerate(bom_data):
        bom.rows[i+1].cells[0].text = comp
        bom.rows[i+1].cells[1].text = part
        bom.rows[i+1].cells[2].text = qty
        bom.rows[i+1].cells[3].text = cost

    doc.add_paragraph()
    doc.add_paragraph('Estimated Prototype BOM: $120-150 per unit')
    doc.add_paragraph('Estimated Production BOM (1,000 units): $45-65 per unit')

    doc.add_heading('Software Stack', level=2)
    sw_items = [
        'Embedded Firmware: ESP-IDF or Arduino framework with FreeRTOS',
        'Mobile Application: React Native (iOS/Android) with Firebase',
        'Cloud Backend: AWS (Lambda, DynamoDB, IoT Core) - HIPAA compliant',
        'Security: TLS 1.3 encryption, secure boot, OTA firmware updates'
    ]
    for item in sw_items:
        doc.add_paragraph(item, style='List Bullet')

    # Phase 4
    doc.add_heading('Phase 4: Prototyping Strategy', level=1)

    proto_table = doc.add_table(rows=4, cols=4)
    proto_table.style = 'Table Grid'
    proto_headers = ['Phase', 'Scope', 'Timeline', 'Budget']
    for i, h in enumerate(proto_headers):
        proto_table.rows[0].cells[i].text = h
        proto_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    proto_data = [
        ('Alpha', 'Mechanism proof-of-concept, basic firmware', '8-12 weeks', '$2,000-3,000'),
        ('Beta', 'Full features, app integration, user testing', '12-16 weeks', '$10,000-15,000'),
        ('Pilot', 'Pre-production for clinical validation', '8-12 weeks', '$25,000-40,000')
    ]
    for i, (phase, scope, time, budget) in enumerate(proto_data):
        proto_table.rows[i+1].cells[0].text = phase
        proto_table.rows[i+1].cells[1].text = scope
        proto_table.rows[i+1].cells[2].text = time
        proto_table.rows[i+1].cells[3].text = budget

    # Phase 5
    doc.add_heading('Phase 5: Manufacturing & Supply Chain', level=1)

    doc.add_heading('Cost Estimates at Scale', level=2)
    cost_table = doc.add_table(rows=4, cols=5)
    cost_table.style = 'Table Grid'
    cost_headers = ['Volume', 'BOM Cost', 'Assembly', 'Total COGS', 'Suggested MSRP']
    for i, h in enumerate(cost_headers):
        cost_table.rows[0].cells[i].text = h
        cost_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    cost_data = [
        ('100 units', '$85', '$35', '$120', '$299'),
        ('1,000 units', '$55', '$20', '$75', '$199'),
        ('10,000 units', '$40', '$12', '$52', '$149')
    ]
    for i, row_data in enumerate(cost_data):
        for j, val in enumerate(row_data):
            cost_table.rows[i+1].cells[j].text = val

    # Phase 6
    doc.add_heading('Phase 6: Regulatory & Compliance', level=1)

    doc.add_heading('FDA Classification', level=2)
    doc.add_paragraph('Expected Classification: Class II Medical Device')
    doc.add_paragraph('Regulatory Pathway: 510(k) Premarket Notification')
    doc.add_paragraph('Estimated 510(k) Timeline: 6-12 months')
    doc.add_paragraph('Estimated 510(k) Cost: $50,000-150,000')

    doc.add_heading('Additional Certifications Required', level=2)
    cert_table = doc.add_table(rows=6, cols=3)
    cert_table.style = 'Table Grid'
    cert_headers = ['Certification', 'Requirement', 'Est. Cost']
    for i, h in enumerate(cert_headers):
        cert_table.rows[0].cells[i].text = h
        cert_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    cert_data = [
        ('FCC Part 15', 'Radio emissions (Wi-Fi, BLE)', '$3,000-5,000'),
        ('CE Mark', 'European market', '$10,000-20,000'),
        ('UL/CSA', 'Electrical safety', '$10,000-15,000'),
        ('HIPAA', 'Data privacy', 'Operational'),
        ('ISO 13485', 'Quality management', '$20,000-50,000')
    ]
    for i, (cert, req, cost) in enumerate(cert_data):
        cert_table.rows[i+1].cells[0].text = cert
        cert_table.rows[i+1].cells[1].text = req
        cert_table.rows[i+1].cells[2].text = cost

    # Phase 7
    doc.add_heading('Phase 7: Business & Legal Structure', level=1)
    doc.add_paragraph('Recommended Entity: Delaware C-Corporation (preferred by investors)')

    doc.add_heading('IP Filings', level=2)
    ip_table = doc.add_table(rows=5, cols=3)
    ip_table.style = 'Table Grid'
    ip_headers = ['Filing', 'Timeline', 'Est. Cost']
    for i, h in enumerate(ip_headers):
        ip_table.rows[0].cells[i].text = h
        ip_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    ip_data = [
        ('Provisional Patent', 'Before public disclosure', '$3,000-5,000'),
        ('Utility Patent', 'Within 12 months of provisional', '$15,000-25,000'),
        ('Design Patent', 'Alongside utility', '$2,000-4,000'),
        ('Trademark', 'Before launch', '$1,000-2,000')
    ]
    for i, (filing, timeline, cost) in enumerate(ip_data):
        ip_table.rows[i+1].cells[0].text = filing
        ip_table.rows[i+1].cells[1].text = timeline
        ip_table.rows[i+1].cells[2].text = cost

    # Phase 8
    doc.add_heading('Phase 8: Go-to-Market Strategy', level=1)

    doc.add_heading('Competitive Landscape', level=2)
    comp_table = doc.add_table(rows=6, cols=4)
    comp_table.style = 'Table Grid'
    comp_headers = ['Competitor', 'Price', 'Key Features', 'Weakness']
    for i, h in enumerate(comp_headers):
        comp_table.rows[0].cells[i].text = h
        comp_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    comp_data = [
        ('Hero Health', '$99 + $30/mo', '10 slots, app, sorting', 'Subscription required'),
        ('MedMinder Maya', '$50/mo', 'Locked dispenser, monitoring', 'No ownership'),
        ('PillPack (Amazon)', 'Pharmacy service', 'Pre-sorted packets', 'Not a device'),
        ('TabTimer', '$50-150', 'Simple timers', 'No smart features'),
        ('Philips', 'Enterprise', 'Hospital-grade', 'Not consumer')
    ]
    for i, (comp, price, features, weakness) in enumerate(comp_data):
        comp_table.rows[i+1].cells[0].text = comp
        comp_table.rows[i+1].cells[1].text = price
        comp_table.rows[i+1].cells[2].text = features
        comp_table.rows[i+1].cells[3].text = weakness

    doc.add_heading('MedSync Differentiation', level=2)
    diff_items = [
        '20 compartments (vs. typical 7-10)',
        'One-time purchase option (vs. subscription-only)',
        'Prescription OCR integration for easy setup',
        'Multi-medication simultaneous dispensing'
    ]
    for item in diff_items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_heading('Pricing Strategy', level=2)
    doc.add_paragraph('Hardware: $149-199 MSRP')
    doc.add_paragraph('Optional Premium Subscription: $9.99/month (advanced analytics, pharmacy integration, cellular backup)')

    # Phase 9
    doc.add_heading('Phase 9: Financial Projections', level=1)

    doc.add_heading('Development Budget', level=2)
    budget_table = doc.add_table(rows=7, cols=2)
    budget_table.style = 'Table Grid'
    budget_data = [
        ('Phase', 'Cost Range'),
        ('Product Development', '$150,000 - 250,000'),
        ('Regulatory (510k)', '$75,000 - 150,000'),
        ('Tooling & Manufacturing Setup', '$100,000 - 200,000'),
        ('Initial Inventory (1,000 units)', '$75,000 - 100,000'),
        ('Marketing & Launch', '$50,000 - 100,000'),
        ('TOTAL TO LAUNCH', '$450,000 - 800,000')
    ]
    for i, (phase, cost) in enumerate(budget_data):
        budget_table.rows[i].cells[0].text = phase
        budget_table.rows[i].cells[1].text = cost
        if i == 0 or i == 6:
            budget_table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            budget_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True

    doc.add_heading('3-Year Revenue Projections', level=2)
    rev_table = doc.add_table(rows=4, cols=5)
    rev_table.style = 'Table Grid'
    rev_headers = ['Year', 'Units Sold', 'Hardware Rev', 'Subscription Rev', 'Total']
    for i, h in enumerate(rev_headers):
        rev_table.rows[0].cells[i].text = h
        rev_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    rev_data = [
        ('1', '2,000', '$350,000', '$60,000', '$410,000'),
        ('2', '8,000', '$1,200,000', '$400,000', '$1,600,000'),
        ('3', '20,000', '$2,800,000', '$1,200,000', '$4,000,000')
    ]
    for i, row_data in enumerate(rev_data):
        for j, val in enumerate(row_data):
            rev_table.rows[i+1].cells[j].text = val

    # Footer
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('Document generated by Product Incubator Skill | February 2026 | Version 1.0')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].italic = True
    footer.runs[0].font.size = Pt(10)

    return doc

if __name__ == '__main__':
    output_dir = Path(__file__).parent
    doc = create_document()
    output_path = output_dir / 'MedSync_Smart_Dispenser_Comprehensive_Plan.docx'
    doc.save(str(output_path))
    print(f'[OK] Comprehensive plan saved to {output_path}')
