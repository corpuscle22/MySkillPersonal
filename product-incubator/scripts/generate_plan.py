"""
Product Incubator - Complete Plan Generator

This script generates the complete product plan including:
1. Workflow diagrams (.drawio format)
2. Schematic images (.png format)
3. Final Word document (.docx format)

All outputs are saved to the outputs/ folder.

Usage:
    python generate_plan.py "Product Name" "Product Description"
    
Example:
    python generate_plan.py "Vertebra Focus Pad" "A focal lumbar massager with remote control"
"""

import argparse
import os
import sys
from pathlib import Path

# Determine paths
SCRIPT_DIR = Path(__file__).parent
SKILL_DIR = SCRIPT_DIR.parent
OUTPUTS_DIR = SKILL_DIR / "outputs"


def get_product_dir(product_name: str) -> Path:
    """Create and return product-specific output directory."""
    safe_name = product_name.lower().replace(' ', '_').replace('-', '_')
    product_dir = OUTPUTS_DIR / safe_name
    product_dir.mkdir(parents=True, exist_ok=True)
    return product_dir, safe_name


def generate_workflow_diagram(product_name: str, output_path: Path):
    """Generate workflow diagram using drawpyo."""
    try:
        import drawpyo
        
        file = drawpyo.File()
        file.file_path = str(output_path)
        page = drawpyo.Page(file=file, name=f"{product_name} Workflow")
        
        steps = [
            "1. Conceptualize & Design",
            "2. IP Check",
            "3. Technical Planning",
            "4. Prototype",
            "5. Manufacturing",
            "6. Business & Legal"
        ]
        
        nodes = []
        y_pos = 50
        for step in steps:
            node = drawpyo.diagram.Object(
                page=page, value=step,
                position=(100, y_pos), width=180, height=50
            )
            nodes.append(node)
            y_pos += 80
        
        for i in range(len(nodes) - 1):
            drawpyo.diagram.Edge(page=page, source=nodes[i], target=nodes[i+1])
        
        file.write()
        print(f"[OK] Workflow diagram saved to {output_path}")
        return True
    except Exception as e:
        print(f"[ERROR] Failed to generate workflow diagram: {e}")
        return False


def generate_system_diagram(product_name: str, output_path: Path):
    """Generate system block diagram using drawpyo."""
    try:
        import drawpyo
        
        file = drawpyo.File()
        file.file_path = str(output_path)
        page = drawpyo.Page(file=file, name=f"{product_name} System")
        
        # Generic system components
        remote = drawpyo.diagram.Object(page=page, value="User Input", position=(50, 150), width=100, height=60)
        controller = drawpyo.diagram.Object(page=page, value="Controller", position=(200, 150), width=120, height=60)
        actuator = drawpyo.diagram.Object(page=page, value="Actuator", position=(370, 150), width=100, height=60)
        output = drawpyo.diagram.Object(page=page, value="Output", position=(520, 150), width=100, height=60)
        
        drawpyo.diagram.Edge(page=page, source=remote, target=controller)
        drawpyo.diagram.Edge(page=page, source=controller, target=actuator)
        drawpyo.diagram.Edge(page=page, source=actuator, target=output)
        
        file.write()
        print(f"[OK] System diagram saved to {output_path}")
        return True
    except Exception as e:
        print(f"[ERROR] Failed to generate system diagram: {e}")
        return False


def generate_schematic_images(product_name: str, images_dir: Path):
    """Generate schematic images using PIL."""
    try:
        from PIL import Image, ImageDraw
        
        # Concept Image
        img = Image.new('RGB', (800, 600), color='white')
        d = ImageDraw.Draw(img)
        
        # Draw a placeholder product shape
        d.rounded_rectangle([200, 200, 600, 400], radius=50, fill='#008080', outline='black', width=2)
        d.text((350, 290), product_name, fill='white')
        d.text((300, 550), "Product Concept Visualization", fill='black')
        
        concept_path = images_dir / f"{product_name.lower().replace(' ', '_')}_concept.png"
        img.save(concept_path)
        print(f"[OK] Concept image saved to {concept_path}")
        
        # System Block Diagram Image
        img2 = Image.new('RGB', (800, 400), color='white')
        d2 = ImageDraw.Draw(img2)
        
        boxes = [(50, 150, 150, 250), (200, 150, 300, 250), (350, 150, 450, 250), (500, 150, 600, 250)]
        labels = ["Input", "Process", "Control", "Output"]
        
        for (box, label) in zip(boxes, labels):
            d2.rectangle(box, fill='#E0E0E0', outline='black', width=2)
            d2.text((box[0]+20, box[1]+45), label, fill='black')
        
        for i in range(len(boxes) - 1):
            d2.line([boxes[i][2], 200, boxes[i+1][0], 200], fill='black', width=2)
        
        system_path = images_dir / f"{product_name.lower().replace(' ', '_')}_system.png"
        img2.save(system_path)
        print(f"[OK] System diagram image saved to {system_path}")
        
        return [concept_path, system_path]
    except Exception as e:
        print(f"[ERROR] Failed to generate schematic images: {e}")
        return []


def generate_word_document(product_name: str, description: str, images: list, documents_dir: Path):
    """Generate the final Word document."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'Product Incubator Plan: {product_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Overview
        doc.add_heading('Overview', level=1)
        doc.add_paragraph(description)
        
        # Phase 1: Conceptualization
        doc.add_heading('Phase 1: Conceptualization & Design', level=1)
        doc.add_paragraph('This section covers the core concept, target audience, and visual design of the product.')
        
        # Add concept image if available
        for img_path in images:
            if 'concept' in str(img_path).lower():
                try:
                    doc.add_picture(str(img_path), width=Inches(5))
                    doc.add_paragraph('Figure: Product Concept Visualization')
                except:
                    doc.add_paragraph(f'[Image: {img_path.name}]')
        
        # Phase 2: IP Check
        doc.add_heading('Phase 2: Intellectual Property (IP) Check', level=1)
        doc.add_paragraph('Preliminary patent search results and novelty assessment.')
        doc.add_paragraph('Disclaimer: This is not legal advice. Consult a patent attorney.')
        
        # Phase 3: Technical Planning
        doc.add_heading('Phase 3: Technical Planning & Workflow', level=1)
        doc.add_paragraph('System architecture and Bill of Materials (BOM).')
        
        # Add system diagram if available
        for img_path in images:
            if 'system' in str(img_path).lower():
                try:
                    doc.add_picture(str(img_path), width=Inches(5))
                    doc.add_paragraph('Figure: System Block Diagram')
                except:
                    doc.add_paragraph(f'[Image: {img_path.name}]')
        
        # Phase 4: Prototyping
        doc.add_heading('Phase 4: Prototyping Strategy', level=1)
        doc.add_paragraph('Fabrication methods, tools, and estimated prototype costs.')
        
        # Phase 5: Manufacturing
        doc.add_heading('Phase 5: Manufacturing & Supply Chain', level=1)
        doc.add_paragraph('Component sourcing and manufacturing partners.')
        
        # Phase 6: Business & Legal
        doc.add_heading('Phase 6: Business & Legal Considerations', level=1)
        doc.add_paragraph('Business entity type, registration, and certifications.')
        
        # Save
        safe_name = product_name.lower().replace(' ', '_')
        output_path = documents_dir / f"{safe_name}_plan.docx"
        doc.save(output_path)
        print(f"[OK] Word document saved to {output_path}")
        return output_path
    except Exception as e:
        print(f"[ERROR] Failed to generate Word document: {e}")
        return None


def main():
    parser = argparse.ArgumentParser(description="Generate complete product plan")
    parser.add_argument("product_name", help="Name of the product")
    parser.add_argument("description", nargs="?", default="A new product concept.", help="Product description")
    
    args = parser.parse_args()
    product_name = args.product_name
    description = args.description
    
    print(f"\n[STARTING] Generating plan for: {product_name}")
    print("=" * 50)
    
    # Create product-specific output folder
    product_dir, safe_name = get_product_dir(product_name)
    
    # Generate diagrams (saved to product folder)
    workflow_path = product_dir / f"{safe_name}_workflow.drawio"
    system_path = product_dir / f"{safe_name}_system.drawio"
    
    generate_workflow_diagram(product_name, workflow_path)
    generate_system_diagram(product_name, system_path)
    
    # Generate images (saved to product folder)
    images = generate_schematic_images(product_name, product_dir)
    
    # Generate Word document (saved to product folder)
    docx_path = generate_word_document(product_name, description, images, product_dir)
    
    print("\n" + "=" * 50)
    print(f"[FOLDER] All outputs saved to: {product_dir}")
    print(f"   - {safe_name}_workflow.drawio")
    print(f"   - {safe_name}_system.drawio")
    print(f"   - {safe_name}_concept.png")
    print(f"   - {safe_name}_system.png")
    print(f"   - {safe_name}_plan.docx")
    print("\n[DONE] Plan generation complete!")


if __name__ == "__main__":
    main()

