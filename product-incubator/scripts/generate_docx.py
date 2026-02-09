from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

def create_proposal_doc(output_path):
    doc = Document()
    
    # Title
    title = doc.add_heading('Product Incubator Plan: The "Vertebra-Focus" Pad', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Overview
    doc.add_heading('Overview', level=1)
    p = doc.add_paragraph('This plan outlines the development of a focal, single-vertebra lumbar massager designed to rest on a bed, providing circular and pressing movements with a soft, finger-like touch.')
    
    # Phase 1
    doc.add_heading('Phase 1: Conceptualization & Design', level=1)
    
    doc.add_heading('Core Concept', level=2)
    p = doc.add_paragraph()
    p.add_run('Product Name: ').bold = True
    p.add_run('Vertebra-Focus Pad (Working Title)\n')
    p.add_run('Form Factor: ').bold = True
    p.add_run('A slim, low-profile pad (approx. 3-4 inches thick) that rests flat on a mattress.\n')
    p.add_run('User Experience: ').bold = True
    p.add_run('The user lies supine (on their back) or on their side. The device is positioned under the lumbar region (specifically L4-L5).\n')
    p.add_run('Key Differentiator: ').bold = True
    p.add_run('Unlike full-back massage mats, this is a focal device with a "Variable-Height" mechanism that allows the user to move the massage nodes up or down to target a specific vertebra using a remote.\n')
    p.add_run('Interface: ').bold = True
    p.add_run('Soft, medical-grade silicone nodes (Shore A 20-30 hardness) to mimic human fingers, avoiding the "hard plastic" feel of typical massagers.')

    # Visual Description with Image
    doc.add_heading('Visual Description', level=2)
    try:
        doc.add_picture('vertebra_concept.png', width=Inches(5))
        doc.add_paragraph('Figure 1: Conceptual Design showing the pad and remote.', style='Caption')
    except:
        doc.add_paragraph('[Image: Vertebra Concept Diagram]')
    
    p = doc.add_paragraph()
    p.add_run('Form Factor: ').bold = True
    p.add_run('A streamlined, pill-shaped pad (approx. 30cm x 15cm x 5cm) with soft, rounded edges to slide easily under the back.\n')
    p.add_run('Color & Finish: ').bold = True
    p.add_run('Calming Teal (Pantone 321C) fabric cover with a smooth, matte white control interface. The silicone nodes are a matching teal but with a soft-touch, skin-like texture.\n')
    p.add_run('Remote: ').bold = True
    p.add_run('Small, pebble-shaped remote in matte white with simple tactile buttons (Up, Down, Knead).')
    
    doc.add_heading('Mechanism Concept', level=2)
    try:
        doc.add_picture('vertebra_mechanism.png', width=Inches(5))
        doc.add_paragraph('Figure 3: Internal Mechanism Schematic (Top View)', style='Caption')
    except:
        doc.add_paragraph('[Mechanism Diagram Placeholder]')

    p = doc.add_paragraph()
    p.add_run('Movement 1 (Massage): ').bold = True
    p.add_run('Dual eccentrically mounted nodes driven by a high-torque DC motor to create circular/kneading motion.\n')
    p.add_run('Movement 2 (Positioning): ').bold = True
    p.add_run('The entire massage module sits on a mini linear track (lead screw or belt drive) approx. 15cm long. A separate motor moves the module up/down the track based on remote input ("Move Up" / "Move Down").\n')
    p.add_run('Pressure: ').bold = True
    p.add_run('Achieving "pressing motion" can be done by the natural weight of the user pressing down on the nodes, or by a simple cam mechanism that pushes the nodes outward (Z-axis) slightly.')

    # Phase 2
    doc.add_heading('Phase 2: Intellectual Property (IP) Check', level=1)
    
    doc.add_heading('Preliminary Search Results', level=2)
    p = doc.add_paragraph()
    p.add_run('Existing Art: ').bold = True
    p.add_run('"Acupressure mats" and "Shiatsu massage pillows" exist but typically lack the remote-controlled positional adjustment (moving the mechanism itself up/down a track) in a small pad format.\n')
    p.add_run('Differentiator: ').bold = True
    p.add_run('The combination of "small/portable pad" + "linear track for focal targeting" appear to be a novel integration space.\n')
    p.add_run('Patent Caution: ').bold = True
    p.add_run('Beware of "SL-Track" patents (massage chairs). Ensure your mechanism is distinct (e.g., a simple flat linear actuator, not a complex 3D curve).')
    
    p = doc.add_paragraph('Disclaimer: I am an AI, not a patent attorney. A professional patent search (approx. cost $1,000 - $3,000) is highly recommended before commercialization.')
    p.runs[0].font.italic = True

    # Phase 3
    doc.add_heading('Phase 3: Technical Planning', level=1)
    
    doc.add_heading('System Diagram', level=2)
    try:
        doc.add_picture('vertebra_system.png', width=Inches(6))
        doc.add_paragraph('Figure 2: System Block Diagram (Remote to Motors)', style='Caption')
    except:
        doc.add_paragraph('[System Diagram Placeholder]')

    doc.add_heading('Bill of Materials (BOM) - Prototype', level=2)
    items = [
        'Microcontroller: ESP32 or Arduino Nano (~$5-10)',
        'Motors: 1x 12V High-Torque DC Geared Motor (for kneading) (~$15), 1x Linear Actuator or Stepper Motor with Lead Screw (15cm travel) (~$30)',
        'Motor Drivers: 2x L298N or similar (~$5)',
        'Remote: 433MHz RF Remote & Receiver Kit (~$5-8)',
        'Power: 12V 2A Power Adapter (~$10)',
        'Housing: 3D Printed PLA/PETG shell + Silicone cast nodes (~$20 materials)',
        'Cushioning: Memory foam overlay + Fabric cover (~$10)'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    # Phase 4
    doc.add_heading('Phase 4: Prototyping Strategy', level=1)
    
    doc.add_heading('Steps', level=2)
    steps = [
        'Mechanical Proof: Build the "Track" first. Mount a motor on a sliding rail. Ensure it can move 20lbs without stalling.',
        'Node Casting: 3D print a mold for the massage nodes. Pour skin-safe silicone to get the "soft finger" feel.',
        'Electronics: Breadboard the RF receiver to control the two motors.',
        'Assembly: Encase in a soft foam block to test comfort on a bed.'
    ]
    for step in steps:
        doc.add_paragraph(step, style='List Number')

    doc.add_heading('Estimated Prototype Cost', level=2)
    p = doc.add_paragraph()
    p.add_run('DIY (Self-built): ').bold = True
    p.add_run('$200 - $400 for parts and materials.\n')
    p.add_run('Professional Prototype: ').bold = True
    p.add_run('$10,000 - $30,000 (Hiring an engineering firm to design custom PCB, CAD, and mechanism).')

    # Phase 5
    doc.add_heading('Phase 5: Manufacturing & Supply Chain', level=1)
    
    doc.add_heading('Sourcing', level=2)
    sources = [
        'Motors: Shenzen-based suppliers on Alibaba.',
        'Silicone Nodes: Custom injection molding (requires tooling).',
        'Assembly: Contract Manufacturers (CM) in China.'
    ]
    for source in sources:
        doc.add_paragraph(source, style='List Bullet')

    doc.add_heading('Product Cost Analysis (at 1,000 units)', level=2)
    p = doc.add_paragraph()
    p.add_run('Target BOM Cost: ').bold = True
    p.add_run('$25 - $40 per unit.\n')
    p.add_run('Retail Price Target: ').bold = True
    p.add_run('$129 - $159 (typical 4x multiple).')

    # Phase 6
    doc.add_heading('Phase 6: Business & Legal', level=1)
    
    doc.add_heading('Regulatory (Crucial!)', level=2)
    p = doc.add_paragraph()
    p.add_run('FDA: ').bold = True
    p.add_run('Class I, Product Code ISA. Generally 510(k) Exempt. Must Register and List annually.\n')
    p.add_run('FCC: ').bold = True
    p.add_run('Remote needs certification. Cost ~$3k-$5k.\n')
    p.add_run('Safety: ').bold = True
    p.add_run('UL 1647 recommended for retail.')

    doc.add_heading('Business Setup', level=2)
    setup = [
        'Entity: Form an LLC (Limited Liability Company) for liability protection.',
        'Insurance: get Product Liability Insurance covering bodily injury.'
    ]
    for s in setup:
        doc.add_paragraph(s, style='List Number')

    # Save
    doc.save(output_path)
    print(f"Document saved to {output_path}")

if __name__ == "__main__":
    create_proposal_doc('Vertebra_Focus_Plan.docx')
